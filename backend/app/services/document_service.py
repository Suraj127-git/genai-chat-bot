from datetime import datetime
from typing import List, Optional, BinaryIO
from fastapi import HTTPException, status, UploadFile
from motor.motor_asyncio import AsyncIOMotorGridFSBucket
from bson import ObjectId
import io
import PyPDF2
import pdfplumber
from docx import Document as DocxDocument
from app.models.document import Document, DocumentCreate, DocumentInDB, DocumentDetail, DocumentType
from app.core.config import settings
from app.db.chromadb import ChromaDBManager
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.embeddings import HuggingFaceEmbeddings
import logging

logger = logging.getLogger(__name__)


class DocumentService:
    """Service for document upload, processing, and management."""
    
    def __init__(self, db):
        self.db = db
        self.documents_collection = db["documents"]
        self.fs_bucket = AsyncIOMotorGridFSBucket(db)
        self.chroma_client = ChromaDBManager.get_client()
        
        # Initialize text splitter for chunking
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            length_function=len,
        )
        
        # Initialize embeddings model
        self.embeddings = HuggingFaceEmbeddings(
            model_name="sentence-transformers/all-MiniLM-L6-v2"
        )
    
    async def upload_and_process(self, file: UploadFile, user_id: str) -> Document:
        """
        Upload and process a document.
        
        Args:
            file: Uploaded file
            user_id: ID of the user uploading the document
            
        Returns:
            Created document object
            
        Raises:
            HTTPException: If file type is invalid or processing fails
        """
        # Validate file type
        file_ext = file.filename.split('.')[-1].lower()
        if file_ext not in settings.allowed_extensions:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=f"File type not allowed. Allowed types: {', '.join(settings.allowed_extensions)}"
            )
        
        # Validate file size
        file_content = await file.read()
        file_size = len(file_content)
        max_size_bytes = settings.MAX_UPLOAD_SIZE_MB * 1024 * 1024
        
        if file_size > max_size_bytes:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=f"File too large. Maximum size: {settings.MAX_UPLOAD_SIZE_MB}MB"
            )
        
        try:
            # Store file in GridFS
            file_id = await self.fs_bucket.upload_from_stream(
                file.filename,
                io.BytesIO(file_content),
                metadata={"user_id": user_id, "content_type": file.content_type}
            )
            
            # Create document record
            doc_dict = {
                "user_id": user_id,
                "filename": file.filename,
                "file_type": file_ext,
                "file_size": file_size,
                "file_path": str(file_id),
                "upload_date": datetime.utcnow(),
                "processed": False,
                "chunk_ids": []
            }
            
            result = await self.documents_collection.insert_one(doc_dict)
            document_id = str(result.inserted_id)
            
            # Process document asynchronously (extract text and create embeddings)
            try:
                await self._process_document(document_id, file_content, file_ext, user_id)
            except Exception as e:
                logger.error(f"Error processing document {document_id}: {e}")
                await self.documents_collection.update_one(
                    {"_id": ObjectId(document_id)},
                    {"$set": {"processing_error": str(e)}}
                )
            
            return Document(
                id=document_id,
                filename=file.filename,
                file_type=DocumentType(file_ext),
                file_size=file_size,
                upload_date=doc_dict["upload_date"],
                processed=False
            )
            
        except Exception as e:
            logger.error(f"Error uploading document: {e}")
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail="Failed to upload document"
            )
    
    async def _process_document(self, document_id: str, file_content: bytes, file_type: str, user_id: str):
        """
        Extract text from document and create vector embeddings.
        
        Args:
            document_id: Document ID
            file_content: Raw file bytes
            file_type: File extension
            user_id: User ID
        """
        # Extract text based on file type
        if file_type == "pdf":
            extracted_text = self._extract_text_from_pdf(file_content)
        elif file_type in ["txt"]:
            extracted_text = file_content.decode('utf-8', errors='ignore')
        elif file_type in ["doc", "docx"]:
            extracted_text = self._extract_text_from_docx(file_content)
        else:
            extracted_text = ""
        
        if not extracted_text:
            raise ValueError("No text could be extracted from document")
        
        # Split text into chunks
        chunks = self.text_splitter.split_text(extracted_text)
        
        # Create embeddings and store in ChromaDB
        collection_name = f"user_{user_id}_documents"
        collection = self.chroma_client.get_or_create_collection(name=collection_name)
        
        # Generate chunk IDs
        chunk_ids = [f"{document_id}_chunk_{i}" for i in range(len(chunks))]
        
        # Add to ChromaDB
        collection.add(
            documents=chunks,
            ids=chunk_ids,
            metadatas=[
                {
                    "document_id": document_id,
                    "chunk_index": i,
                    "total_chunks": len(chunks)
                }
                for i in range(len(chunks))
            ]
        )
        
        # Update document record
        await self.documents_collection.update_one(
            {"_id": ObjectId(document_id)},
            {
                "$set": {
                    "extracted_text": extracted_text[:5000],  # Store first 5000 chars
                    "chunk_ids": chunk_ids,
                    "processed": True
                }
            }
        )
        
        logger.info(f"Processed document {document_id}: {len(chunks)} chunks created")
    
    def _extract_text_from_pdf(self, file_content: bytes) -> str:
        """Extract text from PDF using pdfplumber (more accurate than PyPDF2)."""
        try:
            with pdfplumber.open(io.BytesIO(file_content)) as pdf:
                text_parts = []
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        text_parts.append(text)
                return "\n\n".join(text_parts)
        except Exception as e:
            logger.warning(f"pdfplumber failed, trying PyPDF2: {e}")
            # Fallback to PyPDF2
            try:
                pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
                text_parts = []
                for page in pdf_reader.pages:
                    text = page.extract_text()
                    if text:
                        text_parts.append(text)
                return "\n\n".join(text_parts)
            except Exception as e2:
                logger.error(f"Both PDF extractors failed: {e2}")
                raise ValueError("Could not extract text from PDF")
    
    def _extract_text_from_docx(self, file_content: bytes) -> str:
        """Extract text from DOCX file."""
        try:
            doc = DocxDocument(io.BytesIO(file_content))
            return "\n\n".join([paragraph.text for paragraph in doc.paragraphs if paragraph.text])
        except Exception as e:
            logger.error(f"Error extracting text from DOCX: {e}")
            raise ValueError("Could not extract text from DOCX")
    
    async def get_user_documents(self, user_id: str) -> List[Document]:
        """
        Get all documents for a user.
        
        Args:
            user_id: User ID
            
        Returns:
            List of user documents
        """
        cursor = self.documents_collection.find({"user_id": user_id}).sort("upload_date", -1)
        documents = []
        
        async for doc in cursor:
            documents.append(Document(
                id=str(doc["_id"]),
                filename=doc["filename"],
                file_type=DocumentType(doc["file_type"]),
                file_size=doc["file_size"],
                upload_date=doc["upload_date"],
                processed=doc.get("processed", False)
            ))
        
        return documents
    
    async def get_document(self, document_id: str, user_id: str) -> DocumentDetail:
        """
        Get detailed document information.
        
        Args:
            document_id: Document ID
            user_id: User ID (for authorization)
            
        Returns:
            Detailed document object
            
        Raises:
            HTTPException: If document not found or unauthorized
        """
        doc = await self.documents_collection.find_one({"_id": ObjectId(document_id)})
        
        if not doc:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="Document not found"
            )
        
        if doc["user_id"] != user_id:
            raise HTTPException(
                status_code=status.HTTP_403_FORBIDDEN,
                detail="Not authorized to access this document"
            )
        
        return DocumentDetail(
            id=str(doc["_id"]),
            filename=doc["filename"],
            file_type=DocumentType(doc["file_type"]),
            file_size=doc["file_size"],
            upload_date=doc["upload_date"],
            processed=doc.get("processed", False),
            extracted_text=doc.get("extracted_text"),
            chunk_count=len(doc.get("chunk_ids", []))
        )
    
    async def delete_document(self, document_id: str, user_id: str) -> dict:
        """
        Delete a document and its associated data.
        
        Args:
            document_id: Document ID
            user_id: User ID (for authorization)
            
        Returns:
            Success message
            
        Raises:
            HTTPException: If document not found or unauthorized
        """
        doc = await self.documents_collection.find_one({"_id": ObjectId(document_id)})
        
        if not doc:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="Document not found"
            )
        
        if doc["user_id"] != user_id:
            raise HTTPException(
                status_code=status.HTTP_403_FORBIDDEN,
                detail="Not authorized to delete this document"
            )
        
        # Delete from GridFS
        try:
            await self.fs_bucket.delete(ObjectId(doc["file_path"]))
        except Exception as e:
            logger.warning(f"Could not delete GridFS file: {e}")
        
        # Delete from ChromaDB
        if doc.get("chunk_ids"):
            try:
                collection_name = f"user_{user_id}_documents"
                collection = self.chroma_client.get_or_create_collection(name=collection_name)
                collection.delete(ids=doc["chunk_ids"])
            except Exception as e:
                logger.warning(f"Could not delete ChromaDB chunks: {e}")
        
        # Delete document record
        await self.documents_collection.delete_one({"_id": ObjectId(document_id)})
        
        return {"message": "Document deleted successfully"}
